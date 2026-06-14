import os
import sys
from datetime import date, timedelta


def generate_summary(data, output_path=None):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Missing dependency. Run: pip install python-docx")
        sys.exit(1)

    today = date.today()
    cutoff = today + timedelta(days=7)

    overdue = [
        a for a in data["assignments"]
        if not a["done"] and date.fromisoformat(a["due"]) < today
    ]
    this_week = [
        a for a in data["assignments"]
        if not a["done"] and today <= date.fromisoformat(a["due"]) <= cutoff
    ]
    completed = [a for a in data["assignments"] if a["done"]]
    later = [
        a for a in data["assignments"]
        if not a["done"] and date.fromisoformat(a["due"]) > cutoff
    ]

    doc = Document()

    title = doc.add_heading("DueList Weekly Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Generated {today.strftime('%B %d, %Y')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    def add_section(heading, assignments, color):
        doc.add_heading(heading, level=1)
        if not assignments:
            p = doc.add_paragraph("None")
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            p.runs[0].italic = True
        else:
            for a in sorted(assignments, key=lambda x: x["due"]):
                due = date.fromisoformat(a["due"])
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(a["title"])
                run.bold = True
                run.font.color.rgb = color
                p.add_run(f"  —  {a['class']}  |  Due {due.strftime('%B %d, %Y')}")

    add_section("Overdue", overdue, RGBColor(0xFF, 0x6B, 0x6B))
    add_section("Due This Week", this_week, RGBColor(0xE6, 0x9A, 0x00))
    add_section("Coming Up Later", later, RGBColor(0x33, 0x33, 0x33))
    add_section("Completed", completed, RGBColor(0x2E, 0x9E, 0x4F))

    if output_path is None:
        output_path = os.path.join(
            os.path.dirname(__file__),
            f"DueList_Summary_{today.isoformat()}.docx",
        )

    doc.save(output_path)
    return output_path
